import google.generativeai as genai
import tkinter as tk
from tkinter import filedialog, messagebox
from pathlib import Path
import io
import os
from docx import Document
from PIL import Image, ImageTk
import fitz  # PyMuPDF for PDF processing
from dotenv import load_dotenv

# Load environment variables
load_dotenv()
api_key = os.getenv('GOOGLE_API_KEY')
genai.configure(api_key=api_key)

def initialize_model():
    generation_config = {"temperature": 0.9}
    return genai.GenerativeModel("gemini-1.5-flash", generation_config=generation_config)

def generate_content(model, image_path, prompt):
    image_part = {"mime_type": "image/jpeg", "data": image_path.getvalue()}
    prompt_parts = [prompt, image_part]
    response = model.generate_content(prompt_parts)
    
    if response.candidates:
        candidate = response.candidates[0]
        if candidate.content and candidate.content.parts:
            return candidate.content.parts[0].text or "No valid content generated."
    return "No response received."

def create_word_file(results):
    doc = Document()
    doc.add_heading('Generated Descriptions', 0)
    for description in results:
        doc.add_paragraph(description)
    word_file = "medical_insights.docx"
    doc.save(word_file)
    messagebox.showinfo("Success", f"File saved as {word_file}")

def pdf_to_images(pdf_path):
    images = []
    pdf_document = fitz.open(pdf_path)
    for page in pdf_document:
        pix = page.get_pixmap()
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        img_io = io.BytesIO()
        img.save(img_io, format="JPEG")
        img_io.seek(0)
        images.append(img_io)
    return images

def upload_file():
    file_path = filedialog.askopenfilename(filetypes=[("Image Files", "*.jpg;*.jpeg;*.png"), ("PDF Files", "*.pdf")])
    if file_path:
        process_file(file_path)

def process_file(file_path):
    results.clear()
    model = initialize_model()
    prompt = prompt_text.get("1.0", tk.END).strip() or predefined_prompt
    
    if file_path.lower().endswith(".pdf"):
        images = pdf_to_images(file_path)
        for img in images:
            result = generate_content(model, img, prompt)
            results.append(result)
    else:
        with open(file_path, "rb") as img_file:
            img_io = io.BytesIO(img_file.read())
        result = generate_content(model, img_io, prompt)
        results.append(result)
        display_image(file_path)
    
    result_text.delete("1.0", tk.END)
    result_text.insert(tk.END, "\n".join(results))

def save_results():
    if results:
        create_word_file(results)
    else:
        messagebox.showwarning("Warning", "No results to save.")

def display_image(file_path):
    img = Image.open(file_path)
    img.thumbnail((300, 300))
    img = ImageTk.PhotoImage(img)
    img_label.configure(image=img)
    img_label.image = img

# GUI Setup
root = tk.Tk()
root.title("Med-Vision AI Analyzer")
root.geometry("600x500")
predefined_prompt = "Analyze the medical scan and provide insights."
results = []

upload_btn = tk.Button(root, text="Upload File", command=upload_file)
upload_btn.pack(pady=10)

prompt_label = tk.Label(root, text="Enter Prompt (optional):")
prompt_label.pack()

prompt_text = tk.Text(root, height=3, width=50)
prompt_text.pack()

img_label = tk.Label(root)
img_label.pack()

result_text = tk.Text(root, height=10, width=70)
result_text.pack()

save_btn = tk.Button(root, text="Save Results", command=save_results)
save_btn.pack(pady=10)

root.mainloop()
